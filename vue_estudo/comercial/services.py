import re
from pathlib import Path

from django.conf import settings
from django.core.files import File
from django.utils import timezone


PLACEHOLDER_RE = re.compile(r"{{\s*([a-zA-Z0-9_]+)\s*}}")


def extract_placeholders(file_path):
    path = Path(file_path)
    suffix = path.suffix.lower()
    text = ""

    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError:
            return []
        document = Document(str(path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                text += "\n" + " ".join(cell.text for cell in row.cells)
    elif suffix == ".pdf":
        try:
            import fitz
        except ImportError:
            return []
        with fitz.open(str(path)) as doc:
            text = "\n".join(page.get_text() for page in doc)
    else:
        text = path.read_text(encoding="utf-8", errors="ignore")

    return sorted(set(PLACEHOLDER_RE.findall(text)))


def render_document_with_values(contrato):
    original = Path(contrato.documento_modelo.path)
    output_dir = Path(settings.MEDIA_ROOT) / "documentos" / "finais"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_name = f"contrato_{contrato.pk}_{timezone.now().strftime('%Y%m%d%H%M%S')}{original.suffix}"
    output_path = output_dir / output_name

    suffix = original.suffix.lower()
    values = contrato.valores_preenchidos or {}

    if suffix == ".docx":
        _render_docx(original, output_path, values)
    elif suffix == ".pdf":
        _render_pdf(original, output_path, values)
    else:
        text = original.read_text(encoding="utf-8", errors="ignore")
        output_path.write_text(_replace_text(text, values), encoding="utf-8")

    with output_path.open("rb") as final_file:
        contrato.documento_final.save(output_name, File(final_file), save=True)

    return contrato.documento_final


def _replace_text(text, values):
    for key, value in values.items():
        text = re.sub(r"{{\s*" + re.escape(key) + r"\s*}}", str(value), text)
    return text


def _render_docx(original, output_path, values):
    from docx import Document

    document = Document(str(original))

    for paragraph in document.paragraphs:
        _replace_in_runs(paragraph.runs, values)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_runs(paragraph.runs, values)

    document.save(str(output_path))


def _replace_in_runs(runs, values):
    for run in runs:
        run.text = _replace_text(run.text, values)


def _render_pdf(original, output_path, values):
    import fitz

    with fitz.open(str(original)) as doc:
        for page in doc:
            for key, value in values.items():
                token = "{{" + key + "}}"
                areas = page.search_for(token)
                for rect in areas:
                    page.add_redact_annot(rect, fill=(1, 1, 1))
                    page.apply_redactions()
                    page.insert_text(
                        rect.bottom_left,
                        str(value),
                        fontsize=max(rect.height * 0.8, 8),
                        color=(0, 0, 0),
                    )
        doc.save(str(output_path))


def gerar_pdf_orcamento(orcamento):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas

    output_dir = Path(settings.MEDIA_ROOT) / "orcamentos"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"orcamento_{orcamento.pk}.pdf"

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    preto = colors.HexColor("#111111")
    cinza = colors.HexColor("#9B9B9B")
    cinza_claro = colors.HexColor("#DCDCDC")
    linha = colors.HexColor("#CFCFCF")

    empresa = _orcamento_empresa_data(orcamento)
    logo = empresa["logo"]
    left = 1.15 * cm
    right = width - 1.15 * cm
    top = height - 1.0 * cm

    if logo:
        c.drawImage(
            logo,
            left,
            top - 1.55 * cm,
            width=1.9 * cm,
            height=1.45 * cm,
            preserveAspectRatio=True,
            mask="auto",
        )

    c.setFillColor(preto)
    c.setFont("Helvetica-Bold", 12)
    c.drawString(left + 2.5 * cm, top - 0.25 * cm, empresa["nome"])
    c.setFont("Helvetica-Bold", 7)
    c.drawString(left + 2.5 * cm, top - 0.55 * cm, empresa["cnpj"])
    c.drawString(left + 2.5 * cm, top - 0.85 * cm, empresa["endereco_linha1"])
    c.drawString(left + 2.5 * cm, top - 1.15 * cm, empresa["endereco_linha2"])

    c.setFont("Helvetica-Bold", 7)
    c.drawRightString(right, top - 0.25 * cm, empresa["email"])
    c.drawRightString(right, top - 0.55 * cm, f"Contato: {empresa['telefone']}")

    y = top - 2.25 * cm
    c.setStrokeColor(linha)
    c.setLineWidth(0.4)
    c.line(left, y, right, y)
    y -= 0.45 * cm

    c.setFillColor(preto)
    c.setFont("Helvetica", 10)
    c.drawString(left, y + 0.15 * cm, "Dados do Cliente")
    y -= 0.6 * cm
    c.line(left, y + 0.25 * cm, right, y + 0.25 * cm)
    c.setFont("Helvetica-Bold", 8)
    c.drawString(left, y, orcamento.cliente.nome_completo)
    y -= 0.32 * cm
    c.drawString(left, y, f"CPF/CNPJ: {orcamento.cliente.cpf}")
    y -= 0.32 * cm
    c.drawString(left, y, orcamento.cliente.endereco_residencial or "-")
    y -= 0.32 * cm
    c.drawString(left, y, orcamento.cliente.email or "")
    c.setFont("Helvetica", 9)
    c.drawRightString(right, y + 0.2 * cm, f"Data: {orcamento.criado_em:%d/%m/%Y}")
    y -= 0.6 * cm

    c.setFillColor(cinza)
    c.rect(left, y - 0.08 * cm, right - left, 0.5 * cm, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 10)
    numero = f"ORÇAMENTO Nº {orcamento.pk:04d}-{str(orcamento.criado_em.year)[-2:]}"
    c.drawCentredString(width / 2, y + 0.08 * cm, numero)
    y -= 0.75 * cm

    c.setFillColor(preto)
    c.setFont("Helvetica", 10)
    c.drawString(left, y, "Serviços")
    y -= 0.35 * cm
    c.line(left, y, right, y)
    y -= 0.38 * cm

    c.setFillColor(cinza_claro)
    c.rect(left, y - 0.08 * cm, right - left, 0.42 * cm, fill=1, stroke=0)
    c.setFillColor(preto)
    c.setFont("Helvetica-Bold", 7)
    columns = {
        "nome": left + 0.05 * cm,
        "quantidade": right - 7.45 * cm,
        "unidade": right - 5.8 * cm,
        "valor_unitario": right - 3.6 * cm,
        "valor_total": right,
    }
    c.drawString(columns["nome"], y + 0.05 * cm, "Nome")
    c.drawRightString(columns["quantidade"], y + 0.05 * cm, "Quantidade")
    c.drawRightString(columns["unidade"], y + 0.05 * cm, "Unidade")
    c.drawRightString(columns["valor_unitario"], y + 0.05 * cm, "Valor Unitário")
    c.drawRightString(columns["valor_total"], y + 0.05 * cm, "Valor Total")
    y -= 0.4 * cm

    itens = list(orcamento.itens_produto.select_related("produto")) + list(orcamento.itens_servico.select_related("servico"))
    if itens:
        c.setFont("Helvetica-Bold", 6.5)
        for item in itens:
            label = _orcamento_item_label(item)
            unidade = getattr(getattr(item, "produto", None), "unidade_medida", "un")
            c.drawString(columns["nome"], y, label[:55])
            c.drawRightString(columns["quantidade"], y, _format_quantity(item.quantidade))
            c.drawRightString(columns["unidade"], y, unidade or "un")
            c.drawRightString(columns["valor_unitario"], y, _format_brl(item.valor_unitario))
            c.drawRightString(columns["valor_total"], y, _format_brl(item.total))
            y -= 0.32 * cm
    else:
        c.setFont("Helvetica", 8)
        c.drawString(columns["nome"], y, "Nenhum item informado.")
        y -= 0.38 * cm

    y -= 0.35 * cm
    c.setFont("Helvetica-Bold", 7)
    c.drawRightString(right - 2.1 * cm, y, "Total Serviços")
    c.drawRightString(right, y, _format_brl(orcamento.subtotal_produtos + orcamento.subtotal_servicos))
    y -= 1.15 * cm
    c.drawRightString(right - 2.1 * cm, y, "Subtotal")
    c.drawRightString(right, y, _format_brl(orcamento.subtotal_produtos + orcamento.subtotal_servicos))
    y -= 0.38 * cm
    c.drawRightString(right - 2.1 * cm, y, "Total Orçamento")
    c.drawRightString(right, y, _format_brl(orcamento.valor_total))
    y -= 0.75 * cm

    c.setStrokeColor(linha)
    c.line(left, y + 0.35 * cm, right, y + 0.35 * cm)
    c.setFillColor(preto)
    c.setFont("Helvetica", 10)
    c.drawString(left, y, "Observações")
    y -= 0.55 * cm
    c.setFont("Helvetica-Bold", 7)
    c.drawString(left, y, f"Formas de Pagamento: {orcamento.get_forma_pagamento_display()}")
    y -= 0.3 * cm
    c.drawString(left, y, "Condições de Pagamento: À vista, 3x sem juros")
    y -= 0.55 * cm
    if orcamento.observacoes:
        c.setFont("Helvetica", 6.5)
        for line in str(orcamento.observacoes).splitlines():
            if not line.strip():
                y -= 0.25 * cm
                continue
            c.drawString(left, y, line[:130])
            y -= 0.28 * cm

    y -= 0.4 * cm
    c.setFont("Helvetica", 10)
    c.drawString(left, y, "Fotos")
    y -= 0.35 * cm
    c.line(left, y, right, y)
    y -= 0.35 * cm
    photos = _orcamento_photos(orcamento)
    x = left
    for photo, caption in photos[:3]:
        if Path(photo).exists():
            c.drawImage(photo, x, y - 4.25 * cm, width=3.2 * cm, height=4.1 * cm, preserveAspectRatio=True, mask="auto")
            c.setFillColor(preto)
            c.setFont("Helvetica-Bold", 6.5)
            _draw_wrapped_line(c, caption, x, y - 4.55 * cm, 3.2 * cm)
            x += 3.7 * cm

    c.setStrokeColor(linha)
    c.line(width / 2 - 2.2 * cm, 2.2 * cm, width / 2 + 2.2 * cm, 2.2 * cm)
    c.setFillColor(preto)
    c.setFont("Helvetica-Bold", 7)
    c.drawCentredString(width / 2, 1.8 * cm, empresa["nome"])
    c.drawCentredString(width / 2, 1.55 * cm, empresa["telefone"])

    c.save()
    with output_path.open("rb") as pdf_file:
        orcamento.arquivo_pdf.save(output_path.name, File(pdf_file), save=True)

    return orcamento.arquivo_pdf


def _orcamento_empresa_data(orcamento):
    company = orcamento.company
    fallback_logo = Path(settings.BASE_DIR) / "media/orcamentos/referencia/logo.jpeg"
    if not fallback_logo.exists():
        fallback_logo = Path(settings.BASE_DIR) / "logo2.png"

    def pick(attr, default):
        value = getattr(company, attr, "") if company else ""
        return value or default

    endereco = pick("endereco", "Rua Goiás - Panamericano, 60441-005 - Fortaleza/CE")
    endereco_linha1, endereco_linha2 = _split_address(endereco)
    logo = orcamento.logo.path if orcamento.logo else ""
    if not logo and company and company.logo:
        logo = company.logo.path
    if not logo and fallback_logo.exists():
        logo = str(fallback_logo)

    return {
        "nome": pick("nome_empresa", "Dona do Chopp Ltda"),
        "cnpj": pick("cnpj", "44919343000120"),
        "telefone": pick("telefone", "85981423909"),
        "email": pick("email", "donadochopp@gmail.com"),
        "endereco_linha1": endereco_linha1,
        "endereco_linha2": endereco_linha2,
        "logo": logo,
    }


def _split_address(endereco):
    if " - " in endereco:
        first, rest = endereco.split(" - ", 1)
        return first, rest
    if "," in endereco:
        first, rest = endereco.split(",", 1)
        return first.strip(), rest.strip()
    return endereco, ""


def _format_brl(value):
    formatted = f"{value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    return f"R$ {formatted}"


def _format_quantity(value):
    if value == value.to_integral_value():
        return str(int(value))
    return f"{value:.2f}".replace(".", ",")


def _orcamento_item_label(item):
    if hasattr(item, "produto"):
        produto = item.produto
        partes = [produto.nome]
        if produto.litros and f"{produto.litros}" not in produto.nome:
            partes.append(f"{produto.litros} litros")
        return " - ".join(partes)
    return item.servico.nome


def _orcamento_photos(orcamento):
    photos = []
    for item in orcamento.itens_produto.select_related("produto").prefetch_related("produto__product_images"):
        caption = _orcamento_item_label(item)
        for image in item.produto.product_images.all():
            if image.imagem:
                photos.append((image.imagem.path, caption))
    if photos:
        return photos

    demo_photo = Path(settings.BASE_DIR) / "media/orcamentos/referencia/foto.jpeg"
    if demo_photo.exists():
        first_item = orcamento.itens_produto.select_related("produto").first()
        caption = _orcamento_item_label(first_item) if first_item else "Foto do serviço"
        return [(str(demo_photo), caption)]
    demo_photo = Path(settings.BASE_DIR) / "galeria-chopeira-1.jpeg"
    if demo_photo.exists():
        first_item = orcamento.itens_produto.select_related("produto").first()
        caption = _orcamento_item_label(first_item) if first_item else "Foto do serviço"
        return [(str(demo_photo), caption)]
    return []


def _draw_wrapped_line(c, text, x, y, max_width):
    from reportlab.pdfbase.pdfmetrics import stringWidth
    from reportlab.lib.units import cm

    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if stringWidth(candidate, "Helvetica-Bold", 6.5) <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)

    for line in lines[:2]:
        c.drawString(x, y, line)
        y -= 0.25 * cm
