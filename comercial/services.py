import re
from pathlib import Path

from django.conf import settings
from django.core.files import File
from django.utils import timezone


PLACEHOLDER_RE = re.compile(r"{{\s*([a-zA-Z0-9_]+)\s*}}")

FIXED_CONTRACT_TEMPLATES = {
    "chopeira_eletrica": {
        "filename": "contrato_chopeira_eletrica.docx",
        "download_name": "Contrato Chopeira Eletrica.docx",
    },
    "casamento": {
        "filename": "contrato_casamento.docx",
        "download_name": "Contrato Casamento.docx",
    },
    "contrato_modelo": {
        "filename": "contrato_modelo.docx",
        "download_name": "Contrato Modelo.docx",
    },
}


def fixed_contract_template_path(tipo_modelo):
    template = FIXED_CONTRACT_TEMPLATES.get(tipo_modelo)
    if not template:
        return None
    return Path(settings.BASE_DIR) / "comercial" / "contract_templates" / template["filename"]


def fixed_contract_download_name(tipo_modelo):
    template = FIXED_CONTRACT_TEMPLATES.get(tipo_modelo)
    return template["download_name"] if template else None


def has_fixed_contract_template(tipo_modelo):
    path = fixed_contract_template_path(tipo_modelo)
    return bool(path and path.exists())


def extract_placeholders(file_path):
    path = Path(file_path)
    suffix = path.suffix.lower()
    text = ""

    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError:
            return []
        document = Document(str(path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        for table in document.tables:
            for row in table.rows:
                text += "\n" + " ".join(cell.text for cell in row.cells)
    elif suffix == ".pdf":
        try:
            import fitz
        except ImportError:
            return []
        with fitz.open(str(path)) as doc:
            text = "\n".join(page.get_text() for page in doc)
    else:
        text = path.read_text(encoding="utf-8", errors="ignore")

    return sorted(set(PLACEHOLDER_RE.findall(text)))


def render_document_with_values(contrato):
    original = Path(contrato.documento_modelo.path)
    output_dir = Path(settings.MEDIA_ROOT) / "documentos" / "finais"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_name = f"contrato_{contrato.pk}_{timezone.now().strftime('%Y%m%d%H%M%S')}{original.suffix}"
    output_path = output_dir / output_name

    suffix = original.suffix.lower()
    values = contrato.valores_preenchidos or {}

    if suffix == ".docx":
        _render_docx(original, output_path, values)
    elif suffix == ".pdf":
        _render_pdf(original, output_path, values)
    else:
        text = original.read_text(encoding="utf-8", errors="ignore")
        output_path.write_text(_replace_text(text, values), encoding="utf-8")

    with output_path.open("rb") as final_file:
        contrato.documento_final.save(output_name, File(final_file), save=True)

    return contrato.documento_final


def render_fixed_contract_template(contrato):
    original = fixed_contract_template_path(contrato.tipo_modelo)
    if not original or not original.exists():
        raise FileNotFoundError("Modelo fixo de contrato nao encontrado.")

    output_dir = Path(settings.MEDIA_ROOT) / "documentos" / "finais"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_name = f"contrato_{contrato.pk}_{timezone.now().strftime('%Y%m%d%H%M%S')}{original.suffix}"
    output_path = output_dir / output_name
    values = contrato.valores_preenchidos or {}
    suffix = original.suffix.lower()

    if suffix == ".docx":
        _render_docx(original, output_path, values)
    elif suffix == ".pdf":
        _render_pdf(original, output_path, values)
    else:
        text = original.read_text(encoding="utf-8", errors="ignore")
        output_path.write_text(_replace_text(text, values), encoding="utf-8")

    with output_path.open("rb") as final_file:
        contrato.documento_final.save(output_name, File(final_file), save=True)

    return contrato.documento_final

def _replace_text(text, values):
    for key, value in values.items():
        text = re.sub(r"{{\s*" + re.escape(key) + r"\s*}}", str(value), text)
    return text


def _render_docx(original, output_path, values):
    from docx import Document

    document = Document(str(original))

    for paragraph in document.paragraphs:
        _replace_in_runs(paragraph.runs, values)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_runs(paragraph.runs, values)

    document.save(str(output_path))


def _replace_in_runs(runs, values):
    for run in runs:
        run.text = _replace_text(run.text, values)


def _render_pdf(original, output_path, values):
    import fitz

    with fitz.open(str(original)) as doc:
        for page in doc:
            for key, value in values.items():
                token = "{{" + key + "}}"
                areas = page.search_for(token)
                for rect in areas:
                    page.add_redact_annot(rect, fill=(1, 1, 1))
                    page.apply_redactions()
                    page.insert_text(
                        rect.bottom_left,
                        str(value),
                        fontsize=max(rect.height * 0.8, 8),
                        color=(0, 0, 0),
                    )
        doc.save(str(output_path))


def render_standard_contract(contrato):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    from .models import ConfiguracaoEmpresa

    empresa = ConfiguracaoEmpresa.objects.first()
    nome_empresa = empresa.nome_empresa if empresa else "Dona do Chopp Ltda"
    cnpj_empresa = empresa.cnpj if empresa and empresa.cnpj else ""
    endereco_empresa = empresa.endereco if empresa and empresa.endereco else ""
    email_empresa = empresa.email if empresa and empresa.email else ""
    telefone_empresa = empresa.telefone if empresa and empresa.telefone else ""

    output_dir = Path(settings.MEDIA_ROOT) / "documentos" / "finais"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_name = f"contrato_{contrato.pk}_{timezone.now():%Y%m%d%H%M%S}.docx"
    output_path = output_dir / output_name

    document = Document()
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(
        "CONTRATO DE PRESTAÇÃO DE SERVIÇOS – CHOPEIRA ELÉTRICA"
        if contrato.tipo_modelo == "chopeira_eletrica"
        else "CONTRATO DE PRESTAÇÃO DE SERVIÇOS E COMODATO"
    )
    run.bold = True
    run.font.size = Pt(14)

    cliente = contrato.cliente
    _contract_heading(document, "CONTRATANTE")
    document.add_paragraph(
        f"{cliente.nome_completo}, {cliente.get_tipo_documento_display()} nº "
        f"{cliente.documento_formatado}, telefone: {cliente.celular}, e-mail: "
        f"{cliente.email}, residente em {cliente.endereco_residencial}, doravante "
        "denominado(a) CONTRATANTE."
    )

    _contract_heading(document, "CONTRATADA")
    company_parts = [nome_empresa]
    if cnpj_empresa:
        company_parts.append(f"CNPJ nº {cnpj_empresa}")
    if endereco_empresa:
        company_parts.append(f"com sede em {endereco_empresa}")
    if email_empresa:
        company_parts.append(f"e-mail: {email_empresa}")
    if telefone_empresa:
        company_parts.append(f"telefone: {telefone_empresa}")
    document.add_paragraph(
        ", ".join(company_parts) + ", doravante denominada CONTRATADA."
    )
    document.add_paragraph(
        "As partes firmam o presente contrato de prestação de serviços, regido "
        "pelas cláusulas e condições seguintes."
    )

    _contract_heading(document, "1. OBJETO DO CONTRATO")
    document.add_paragraph(
        "Prestação de serviços e fornecimento dos itens abaixo para o evento "
        "do(a) CONTRATANTE:"
    )
    if contrato.orcamento:
        for item in contrato.orcamento.itens_produto.select_related("produto"):
            document.add_paragraph(
                f"{_contract_quantity(item.quantidade)} × {item.produto.nome}",
                style="List Bullet",
            )
        for item in contrato.orcamento.itens_servico.select_related("servico"):
            document.add_paragraph(
                f"{_contract_quantity(item.quantidade)} × {item.servico.nome}",
                style="List Bullet",
            )
    else:
        document.add_paragraph(
            "Itens conforme proposta comercial aceita pelas partes.", style="List Bullet"
        )

    local = contrato.endereco_evento or "não informado"
    data = contrato.data_evento.strftime("%d/%m/%Y") if contrato.data_evento else "não informada"
    inicio = contrato.horario_inicio.strftime("%H:%M") if contrato.horario_inicio else "não informado"
    fim = contrato.horario_fim.strftime("%H:%M") if contrato.horario_fim else "não informado"
    document.add_paragraph(f"Local do evento: {local}")
    document.add_paragraph(f"Data do evento: {data}")
    document.add_paragraph(f"Horário: {inicio} às {fim}")
    document.add_paragraph(
        "Os equipamentos deverão permanecer no endereço indicado, sendo vedado "
        "seu deslocamento sem autorização prévia da CONTRATADA."
    )

    _contract_heading(document, "2. RESPONSABILIDADE PELOS EQUIPAMENTOS")
    document.add_paragraph(
        "O(A) CONTRATANTE compromete-se a utilizar e conservar adequadamente os "
        "equipamentos, respondendo por danos, perdas ou extravios durante o período "
        "em que estiverem sob sua responsabilidade."
    )
    if not contrato.com_profissional:
        document.add_paragraph(
            "Como o serviço foi contratado sem profissional da CONTRATADA, a "
            "operação segura e a guarda dos equipamentos serão de responsabilidade "
            "integral do(a) CONTRATANTE."
        )

    _contract_heading(document, "3. PROFISSIONAIS E HORAS EXTRAS")
    if contrato.com_profissional:
        document.add_paragraph(
            f"O serviço inclui {contrato.quantidade_profissionais} profissional(is) "
            f"no horário contratado. Cada hora adicional iniciada terá o valor de "
            f"{_contract_money(contrato.valor_hora_extra)}, pago diretamente no "
            "momento da solicitação."
        )
    else:
        document.add_paragraph("Este contrato não inclui profissional da CONTRATADA.")

    _contract_heading(document, "4. SOBRA DO CHOPP E PERMANÊNCIA DOS EQUIPAMENTOS")
    document.add_paragraph(
        f"Havendo sobra de chopp, a chopeira poderá permanecer no mesmo endereço "
        f"por até {contrato.prazo_chopeira_horas} horas, mediante autorização da "
        "CONTRATADA. O conteúdo também poderá ser acondicionado em recipientes "
        "fornecidos pelo(a) CONTRATANTE."
    )
    document.add_paragraph(
        f"O transporte ou a reinstalação em outro endereço, quando autorizados, "
        f"terão taxa de {_contract_money(contrato.taxa_nova_instalacao)}."
    )

    _contract_heading(document, "5. VALOR E FORMA DE PAGAMENTO")
    total = contrato.valor_total
    document.add_paragraph(f"Valor total: {_contract_money(total)}.")
    document.add_paragraph(
        f"Valor pago: {_contract_money(contrato.valor_pago)}"
        + (
            f", em {contrato.data_pagamento:%d/%m/%Y}."
            if contrato.data_pagamento else "."
        )
    )
    document.add_paragraph(
        f"Saldo pendente: {_contract_money(contrato.saldo_pendente)}"
        + (
            f", com vencimento em {contrato.data_vencimento_saldo:%d/%m/%Y}."
            if contrato.data_vencimento_saldo else "."
        )
    )
    if contrato.orcamento:
        document.add_paragraph(
            f"Forma de pagamento: {contrato.orcamento.get_forma_pagamento_display()}."
        )

    _contract_heading(document, "6. CANCELAMENTO")
    document.add_paragraph(
        "a) Em até 7 (sete) dias após a assinatura, será devolvido 100% do valor pago."
    )
    document.add_paragraph(
        "b) Após esse prazo, será devolvido 50% do valor pago, considerando a reserva "
        "da data e os custos operacionais já assumidos."
    )
    document.add_paragraph(
        "c) Se a CONTRATADA cancelar, devolverá 100% do valor recebido, ressalvados "
        "casos fortuitos ou de força maior."
    )

    _contract_heading(document, "7. FORO")
    document.add_paragraph(
        f"As partes elegem o foro da comarca de {contrato.cidade_assinatura}, com "
        "renúncia a qualquer outro, para resolver dúvidas decorrentes deste contrato."
    )
    if contrato.observacoes:
        _contract_heading(document, "OBSERVAÇÕES ADICIONAIS")
        document.add_paragraph(contrato.observacoes)

    signature_date = contrato.data_assinatura_usuario or timezone.localdate()
    document.add_paragraph()
    document.add_paragraph(
        f"{contrato.cidade_assinatura}, {signature_date:%d/%m/%Y}."
    )
    document.add_paragraph()
    signature = document.add_table(rows=2, cols=2)
    signature.cell(0, 0).text = "________________________________"
    signature.cell(0, 1).text = "________________________________"
    signature.cell(1, 0).text = f"CONTRATANTE\n{cliente.nome_completo}"
    signature.cell(1, 1).text = f"CONTRATADA\n{nome_empresa}"
    for row in signature.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(str(output_path))
    with output_path.open("rb") as final_file:
        contrato.documento_final.save(output_name, File(final_file), save=True)
    return contrato.documento_final


def _contract_heading(document, text):
    from docx.shared import Pt

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True


def _contract_money(value):
    text = f"{value:,.2f}".replace(",", "_").replace(".", ",").replace("_", ".")
    return f"R$ {text}"


def _contract_quantity(value):
    text = f"{value:.2f}".replace(".", ",")
    return text.rstrip("0").rstrip(",")


def gerar_pdf_orcamento(orcamento):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.pdfgen import canvas
    from reportlab.platypus import Frame, Paragraph

    output_dir = Path(settings.MEDIA_ROOT) / "orcamentos"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / f"orcamento_{orcamento.pk}.pdf"

    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    amarelo = colors.HexColor("#FFC52F")
    vermelho = colors.HexColor("#D8332F")
    preto = colors.HexColor("#191919")

    from .models import ConfiguracaoEmpresa

    configuracao = ConfiguracaoEmpresa.objects.first()
    logo_field = orcamento.logo or (
        configuracao.logo if configuracao and configuracao.logo else None
    )
    logo = _image_reader(logo_field)
    if not logo:
        fallback_logo = Path(settings.BASE_DIR) / "logo2.png"
        if fallback_logo.exists():
            from reportlab.lib.utils import ImageReader
            logo = ImageReader(str(fallback_logo))

    header_top = height - 1.5 * cm
    if logo:
        c.drawImage(
            logo,
            1.5 * cm,
            height - 4.1 * cm,
            width=2.6 * cm,
            height=2.6 * cm,
            preserveAspectRatio=True,
            anchor="c",
            mask="auto",
        )

    company_x = 4.4 * cm
    company_name = configuracao.nome_empresa if configuracao else "Dona do Chopp"
    c.setFillColor(preto)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(company_x, header_top - 0.25 * cm, company_name)
    c.setFont("Helvetica", 9.5)
    company_y = header_top - 0.75 * cm
    if configuracao and configuracao.cnpj:
        c.drawString(company_x, company_y, configuracao.cnpj)
        company_y -= 0.42 * cm
    if configuracao and configuracao.endereco:
        c.drawString(company_x, company_y, configuracao.endereco[:72])

    c.setFont("Helvetica", 9.5)
    contact_y = header_top - 0.25 * cm
    if configuracao and configuracao.email:
        c.drawRightString(width - 1.5 * cm, contact_y, configuracao.email)
        contact_y -= 0.42 * cm
    if configuracao and configuracao.telefone:
        c.drawRightString(width - 1.5 * cm, contact_y, f"Contato: {configuracao.telefone}")

    section_y = height - 4.65 * cm
    c.setFont("Helvetica-Bold", 14)
    c.drawString(1.5 * cm, section_y, "Dados do Cliente")
    c.setStrokeColor(colors.HexColor("#D8D8D8"))
    c.setLineWidth(1)
    c.line(1.5 * cm, section_y - 0.2 * cm, width - 1.5 * cm, section_y - 0.2 * cm)

    cliente = orcamento.cliente
    client_y = section_y - 0.85 * cm
    c.setFont("Helvetica-Bold", 11)
    c.drawString(1.5 * cm, client_y, cliente.nome_completo)
    c.setFont("Helvetica", 9.5)
    c.drawString(
        1.5 * cm,
        client_y - 0.45 * cm,
        f"{cliente.get_tipo_documento_display()}: {cliente.documento_formatado}",
    )
    c.drawString(1.5 * cm, client_y - 0.9 * cm, cliente.endereco_residencial[:86])
    c.drawString(
        1.5 * cm,
        client_y - 1.35 * cm,
        f"Contato: {cliente.celular} | E-mail: {cliente.email}",
    )

    data_orcamento = timezone.localtime(orcamento.criado_em).strftime("%d/%m/%Y")
    validade = orcamento.validade.strftime("%d/%m/%Y") if orcamento.validade else "Não informada"
    c.setFont("Helvetica", 10)
    c.drawRightString(width - 1.5 * cm, client_y - 0.9 * cm, f"Data: {data_orcamento}")
    c.drawRightString(width - 1.5 * cm, client_y - 1.35 * cm, f"Validade: {validade}")

    title_y = height - 7.75 * cm
    c.setFillColor(preto)
    c.setFont("Helvetica-Bold", 18)
    c.drawString(1.5 * cm, title_y, f"Orçamento #{orcamento.pk}")
    c.setFont("Helvetica", 9.5)
    c.drawRightString(
        width - 1.5 * cm,
        title_y,
        f"Pagamento: {orcamento.get_forma_pagamento_display()} | Responsável: {orcamento.responsavel}",
    )

    y = height - 8.65 * cm
    c.setFillColor(vermelho)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(2 * cm, y, "Produtos")
    y -= 0.7 * cm
    y = _draw_items(
        c,
        y,
        orcamento.itens_produto.all(),
        lambda item: item.produto.nome,
        lambda item: item.produto.imagem,
    )

    y -= 0.4 * cm
    c.setFillColor(vermelho)
    c.setFont("Helvetica-Bold", 16)
    c.drawString(2 * cm, y, "Serviços")
    y -= 0.7 * cm
    _draw_items(
        c,
        y,
        orcamento.itens_servico.all(),
        lambda item: item.servico.nome,
        lambda item: item.servico.imagem,
    )

    c.setFillColor(amarelo)
    c.roundRect(2 * cm, 4 * cm, width - 4 * cm, 1.3 * cm, 0.2 * cm, fill=1, stroke=0)
    c.setFillColor(preto)
    c.setFont("Helvetica-Bold", 18)
    c.drawRightString(width - 2.2 * cm, 4.45 * cm, f"Total: R$ {orcamento.valor_total:.2f}")

    if orcamento.observacoes:
        style = ParagraphStyle("obs", fontName="Helvetica", fontSize=9, leading=11, textColor=preto)
        frame = Frame(2 * cm, 2.6 * cm, width - 4 * cm, 1 * cm, showBoundary=0)
        frame.addFromList([Paragraph(orcamento.observacoes, style)], c)

    c.setStrokeColor(preto)
    c.line(2 * cm, 1.7 * cm, 10 * cm, 1.7 * cm)
    c.setFont("Helvetica", 9)
    c.drawString(2 * cm, 1.3 * cm, f"Assinatura do funcionário: {orcamento.responsavel}")

    c.save()
    with output_path.open("rb") as pdf_file:
        orcamento.arquivo_pdf.save(output_path.name, File(pdf_file), save=True)

    return orcamento.arquivo_pdf

def _image_reader(image_field):
    if not image_field:
        return None

    from io import BytesIO

    from reportlab.lib.utils import ImageReader

    try:
        image_field.open("rb")
        content = image_field.read()
        image_field.close()
        return ImageReader(BytesIO(content))
    except (FileNotFoundError, OSError, ValueError):
        return None


def _draw_items(c, y, items, label_getter, image_getter):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm

    if not items:
        c.setFillColor(colors.HexColor("#666666"))
        c.setFont("Helvetica-Oblique", 9)
        c.drawString(2 * cm, y, "Nenhum item informado.")
        return y - 0.55 * cm

    columns = [2 * cm, 3.25 * cm, 11.2 * cm, 14.2 * cm, 16.6 * cm, 19 * cm]
    c.setFillColor(colors.HexColor("#252525"))
    c.rect(2 * cm, y - 0.55 * cm, 17 * cm, 0.65 * cm, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont("Helvetica-Bold", 8.5)
    c.drawString(columns[0] + 0.12 * cm, y - 0.32 * cm, "Foto")
    c.drawString(columns[1], y - 0.32 * cm, "Item")
    c.drawRightString(columns[3], y - 0.32 * cm, "Qtd.")
    c.drawRightString(columns[4], y - 0.32 * cm, "Valor unit.")
    c.drawRightString(columns[5] - 0.15 * cm, y - 0.32 * cm, "Total")
    y -= 0.72 * cm

    for index, item in enumerate(items):
        row_height = 1.1 * cm
        if y - row_height < 5.5 * cm:
            c.showPage()
            y = A4[1] - 2 * cm
        if index % 2:
            c.setFillColor(colors.HexColor("#F5F5F5"))
            c.rect(2 * cm, y - 0.82 * cm, 17 * cm, row_height, fill=1, stroke=0)

        image = _image_reader(image_getter(item))
        if image:
            c.drawImage(
                image,
                columns[0] + 0.1 * cm,
                y - 0.72 * cm,
                width=0.82 * cm,
                height=0.82 * cm,
                preserveAspectRatio=True,
                anchor="c",
                mask="auto",
            )

        text_y = y - 0.4 * cm
        c.setFillColor(colors.HexColor("#191919"))
        c.setFont("Helvetica", 9)
        c.drawString(columns[1], text_y, label_getter(item)[:48])
        c.drawRightString(columns[3], text_y, _numero_pdf(item.quantidade))
        c.drawRightString(columns[4], text_y, _dinheiro_pdf(item.valor_unitario))
        c.drawRightString(columns[5] - 0.15 * cm, text_y, _dinheiro_pdf(item.total))
        c.setStrokeColor(colors.HexColor("#DDDDDD"))
        c.line(2 * cm, y - 0.83 * cm, 19 * cm, y - 0.83 * cm)
        y -= row_height
    return y


def _dinheiro_pdf(value):
    text = f"{value:,.2f}".replace(",", "_").replace(".", ",").replace("_", ".")
    return f"R$ {text}"


def _numero_pdf(value):
    text = f"{value:,.2f}".replace(",", "_").replace(".", ",").replace("_", ".")
    return text.rstrip("0").rstrip(",")

