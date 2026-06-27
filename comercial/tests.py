from datetime import date, time, timedelta
from decimal import Decimal
from tempfile import TemporaryDirectory

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.urls import reverse

from .forms import ClienteForm, ProdutoForm
from .models import Cliente, Contrato, Orcamento, OrcamentoProduto, Produto
from .services import render_standard_contract


class CommercialFlowTests(TestCase):
    def setUp(self):
        self.user = get_user_model().objects.create_user(
            username="admin",
            password="secret",
            first_name="Igor",
            last_name="Silva",
        )
        self.client.force_login(self.user)
        self.cliente = Cliente.objects.create(
            nome_completo="Cliente Teste",
            tipo_documento="cpf",
            documento="12345678901",
            email="cliente@example.com",
            endereco_residencial="Rua Teste, 10",
            celular="11999999999",
        )
        self.produto = Produto.objects.create(
            nome="Produto sem litros",
            valor=Decimal("100.00"),
            estoque_quantidade=10,
            litros=None,
        )
        self.orcamento = Orcamento.objects.create(
            cliente=self.cliente,
            usuario=self.user,
            forma_pagamento="pix",
        )
        OrcamentoProduto.objects.create(
            orcamento=self.orcamento,
            produto=self.produto,
            quantidade=Decimal("2"),
            valor_unitario=Decimal("100.00"),
        )

    def _contract(self, title, status, event_date):
        return Contrato.objects.create(
            cliente=self.cliente,
            usuario=self.user,
            orcamento=self.orcamento,
            titulo=title,
            status=status,
            data_evento=event_date,
            documento_modelo=SimpleUploadedFile("contrato.txt", b"Contrato"),
        )

    def test_cliente_accepts_cpf_and_cnpj(self):
        cpf_form = ClienteForm(
            data={
                "nome_completo": "Pessoa",
                "tipo_documento": "cpf",
                "documento": "987.654.321-00",
                "email": "pessoa@example.com",
                "endereco_residencial": "Rua A",
                "celular": "11999999999",
            }
        )
        self.assertTrue(cpf_form.is_valid(), cpf_form.errors)
        self.assertEqual(cpf_form.cleaned_data["documento"], "98765432100")

        cnpj_form = ClienteForm(
            data={
                "nome_completo": "Empresa",
                "tipo_documento": "cnpj",
                "documento": "12.345.678/0001-90",
                "email": "empresa@example.com",
                "endereco_residencial": "Rua B",
                "celular": "11888888888",
            }
        )
        self.assertTrue(cnpj_form.is_valid(), cnpj_form.errors)
        self.assertEqual(cnpj_form.cleaned_data["documento"], "12345678000190")

    def test_product_list_has_no_brand_route(self):
        response = self.client.get(reverse("produto_list"))
        self.assertEqual(response.status_code, 200)
        self.assertNotContains(response, "marca_list")
        self.assertNotContains(response, "Marcas")
    def test_produto_litros_is_optional(self):
        form = ProdutoForm(
            data={
                "nome": "Chopeira",
                "disponivel": True,
                "valor": "250.00",
                "estoque_quantidade": 2,
                "litros": "",
                "descricao": "",
            }
        )
        self.assertTrue(form.is_valid(), form.errors)

    def test_orcamento_has_no_status_and_tracks_employee(self):
        self.assertFalse(hasattr(self.orcamento, "status"))
        self.assertEqual(self.orcamento.responsavel, "Igor Silva")

        response = self.client.get(reverse("orcamento_detail", args=[self.orcamento.pk]))
        self.assertContains(response, "Igor Silva")
        self.assertContains(response, "Assinatura do funcionário")

    def test_dashboard_uses_contracts_and_sold_items(self):
        today = date.today()
        self._contract("Executado", "executado", today)
        self._contract("Futuro", "assinado", today + timedelta(days=10))
        self._contract("Cancelado", "cancelado", today + timedelta(days=20))

        response = self.client.get(reverse("dashboard"), {"periodo": "todos"})
        self.assertEqual(response.status_code, 200)
        self.assertEqual(response.context["contratos_executados"], 1)
        self.assertEqual(response.context["contratos_aguardando"], 1)
        self.assertEqual(response.context["contratos_cancelados"], 1)
        self.assertEqual(response.context["total_vendido"], Decimal("200.00"))
        self.assertContains(response, "Produto sem litros")
    def test_standard_contract_generates_docx_from_variable_data(self):
        contrato = Contrato.objects.create(
            cliente=self.cliente,
            usuario=self.user,
            orcamento=self.orcamento,
            titulo="Contrato padrao",
            tipo_modelo="chopeira_eletrica",
            data_evento=date(2026, 7, 10),
            endereco_evento="Rua do Evento, 100",
            horario_inicio=time(18, 0),
            horario_fim=time(22, 0),
            com_profissional=True,
            quantidade_profissionais=2,
            valor_pago=Decimal("50.00"),
        )
        with TemporaryDirectory() as media_root, override_settings(MEDIA_ROOT=media_root):
            arquivo = render_standard_contract(contrato)
            from docx import Document

            document = Document(arquivo.path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("Cliente Teste", text)
            self.assertIn("Produto sem litros", text)
            self.assertIn("Rua do Evento, 100", text)
            self.assertIn("R$ 200,00", text)
            self.assertIn("R$ 150,00", text)

    def test_standard_contract_form_skips_placeholder_screen_and_uses_budget_data(self):
        with TemporaryDirectory() as media_root, override_settings(MEDIA_ROOT=media_root):
            response = self.client.post(
                reverse("contrato_create"),
                data={
                    "tipo_modelo": "chopeira_eletrica",
                    "cliente": self.cliente.pk,
                    "orcamento": self.orcamento.pk,
                    "titulo": "Contrato pelo formulario",
                    "status": "rascunho",
                    "data_evento": "2026-07-10",
                    "endereco_evento": "Rua do Evento, 100",
                    "horario_inicio": "18:00",
                    "horario_fim": "22:00",
                    "com_profissional": "on",
                    "quantidade_profissionais": "1",
                    "valor_hora_extra": "50.00",
                    "valor_pago": "50.00",
                    "data_pagamento": "",
                    "data_vencimento_saldo": "",
                    "prazo_chopeira_horas": "24",
                    "taxa_nova_instalacao": "250.00",
                    "cidade_assinatura": "Fortaleza/CE",
                    "data_assinatura_cliente": "",
                    "data_assinatura_usuario": "",
                    "observacoes": "",
                },
            )

            contrato = Contrato.objects.latest("pk")
            self.assertRedirects(response, reverse("contrato_detail", args=[contrato.pk]))
            self.assertEqual(contrato.placeholders, [])
            self.assertTrue(contrato.documento_final)

            from docx import Document

            document = Document(contrato.documento_final.path)
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)
            self.assertIn("Cliente Teste", text)
            self.assertIn("Produto sem litros", text)
            self.assertIn("Rua do Evento, 100", text)
            self.assertNotIn("Kaua", text)
            self.assertNotIn("{{NOME}}", text)
            self.assertNotIn("{{CONTRATADA}}", text)

